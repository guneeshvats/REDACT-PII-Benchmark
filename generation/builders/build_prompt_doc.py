"""Render generation_prompt_v3.yaml as a styled Word document."""

import json
import yaml
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).parent
SRC  = HERE / "generation_prompt_v3.yaml"
DST  = HERE / "PII_Benchmark_Generation_Prompt_v3.docx"

# ── colors ────────────────────────────────────────────────────────────────────
DARK_BLUE   = RGBColor(0x1F, 0x38, 0x64)
MED_BLUE    = RGBColor(0x2E, 0x75, 0xB6)
ACCENT_BLUE = RGBColor(0xBD, 0xD7, 0xEE)
LIGHT_GREY  = RGBColor(0xF2, 0xF2, 0xF2)
CODE_GREY   = RGBColor(0xF7, 0xF7, 0xF7)
RULE_RED    = RGBColor(0xC0, 0x00, 0x00)
GREEN       = RGBColor(0x37, 0x86, 0x36)
BLACK       = RGBColor(0x10, 0x10, 0x10)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)


def rgb_to_hex(rgb):
    v = tuple(rgb)
    return f"{v[0]:02X}{v[1]:02X}{v[2]:02X}"


def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_to_hex(rgb))
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    sizes = {1: 18, 2: 14, 3: 11}
    colors = {1: DARK_BLUE, 2: MED_BLUE, 3: DARK_BLUE}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10 if level == 2 else 6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(sizes[level])
    r.font.color.rgb = colors[level]
    if level == 3:
        r.italic = True
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLACK
    return p


def add_code_block(doc, text):
    """Render a multi-line code/JSON block as monospace, light-grey background."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_to_hex(CODE_GREY))
    pPr.append(shd)
    r = p.add_run(text)
    r.font.name = "Menlo"
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Menlo")
    rFonts.set(qn("w:hAnsi"), "Menlo")
    rFonts.set(qn("w:cs"), "Menlo")
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return p


def add_text_block(doc, text):
    """Render multi-line prose preserving line breaks."""
    for line in text.rstrip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(line if line else " ")
        r.font.size = Pt(10)
        r.font.color.rgb = BLACK
        # bold the R-rule lines
        stripped = line.lstrip()
        if stripped.startswith("R") and len(stripped) >= 3 and stripped[1].isdigit() and stripped[2:4] in (". ", "  ", "0."):
            r.bold = True
            r.font.color.rgb = DARK_BLUE


def axes_table(doc, axes_dict):
    """Two-column table: axis name | value."""
    table = doc.add_table(rows=len(axes_dict) + 1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header
    h = table.rows[0]
    for i, txt in enumerate(["Axis", "Value"]):
        c = h.cells[i]
        set_cell_bg(c, ACCENT_BLUE)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        run = c.paragraphs[0].add_run(txt)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = DARK_BLUE
    # rows
    for i, (k, v) in enumerate(axes_dict.items(), start=1):
        row = table.rows[i]
        bg = LIGHT_GREY if i % 2 == 0 else WHITE
        for j, txt in enumerate([k, str(v)]):
            c = row.cells[j]
            set_cell_bg(c, bg)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            run = c.paragraphs[0].add_run(txt)
            run.font.size = Pt(9)
            if j == 0:
                run.font.color.rgb = MED_BLUE
                run.bold = True
            else:
                run.font.name = "Menlo"
                rPr = run._element.get_or_add_rPr()
                rFonts = OxmlElement("w:rFonts")
                rFonts.set(qn("w:ascii"), "Menlo")
                rFonts.set(qn("w:hAnsi"), "Menlo")
                rFonts.set(qn("w:cs"), "Menlo")
                rPr.append(rFonts)
                run.font.color.rgb = BLACK
    doc.add_paragraph()
    return table


def pretty_json(json_str):
    """Re-format embedded JSON for readability."""
    s = json_str.strip()
    try:
        obj = json.loads(s)
        return json.dumps(obj, indent=2, ensure_ascii=False)
    except Exception:
        return s


def main():
    data = yaml.safe_load(SRC.read_text(encoding="utf-8"))

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # ── Title ────────────────────────────────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("PII Benchmark — Generation Prompt v3.2")
    tr.bold = True
    tr.font.size = Pt(22)
    tr.font.color.rgb = DARK_BLUE

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run("Specialized prompt for synthetic eval-data generation\n"
                   "51 entities × 9 axes × 25 languages | 16 hard rules | 11 few-shots | Paired-Sweep Mode")
    sr.font.size = Pt(11)
    sr.italic = True
    sr.font.color.rgb = MED_BLUE

    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = d.add_run("Anonymous Institution  ·  May 2026")
    dr.font.size = Pt(9.5)
    dr.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    doc.add_paragraph()

    # ── About section ────────────────────────────────────────────────────────
    add_heading(doc, "About this prompt", level=2)
    add_body(doc,
        "This document is the generation prompt used to produce ground-truth "
        "records for the PII Detection Benchmark (Track 1). It is consumed "
        "by a single LLM call: the model emits the passage and the entity "
        "list; character offsets are computed downstream by deterministic "
        "string alignment (LLMs do not produce offsets — Issues doc §D.12).")
    add_body(doc,
        "It addresses every prompt-addressable gap from the v3 Gap Analysis "
        "(A5 behavioral matrix, B3 negation/hypothetical, B4 partial/obfuscated, "
        "B5 cross-sentence coreference, C3 nested spans, D3/Article-9 sensitivity, "
        "F2 format & checksum validity). Pipeline-, sampler-, and audit-level "
        "gaps are deliberately out of scope.")

    # ── System prompt ────────────────────────────────────────────────────────
    add_heading(doc, "1. System Prompt", level=1)
    add_body(doc,
        "Sent as the system message on every generation call. R1–R15 are "
        "enforced; the output schema is authoritative (overrides any older "
        "few-shot pattern).")
    add_text_block(doc, data["system_prompt"])

    # ── User prompt ──────────────────────────────────────────────────────────
    add_heading(doc, "2. User Prompt Template", level=1)
    add_body(doc,
        "Axis values in {{ }} are interpolated by the sampler (covering "
        "array, not random). Every record gets a fully-resolved 9-axis "
        "configuration plus a behavioral_frame and a prior_passages_summary "
        "for the fixed-axis sweep.")
    add_text_block(doc, data["user_prompt"])

    # ── Few-shots ────────────────────────────────────────────────────────────
    add_heading(doc, "3. Few-Shots", level=1)
    add_body(doc,
        "Eleven examples chosen to anchor the cases v2 prompts got wrong. Each "
        "example shows the axis configuration that produced it (top) and the "
        "exact JSON the model must emit (bottom). Few-shot #1 demonstrates "
        "the FULL v3.1 schema including sensitivity_tier and nesting_parent; "
        "remaining few-shots focus on illustrating their target behavior.")

    fs_titles = {
        0:  "Few-shot #1 — Adjacency=tight / CO-008 email metadata block (multi-mention test, full v3.1 schema)",
        1:  "Few-shot #2 — Adjacency=tight / CO-009 patient header (healthcare)",
        2:  "Few-shot #3 — Hypothetical SSN (disclosed=false)",
        3:  "Few-shot #4 — Partial PII (disclosure_form=partial)",
        4:  "Few-shot #5 — Hard OCR distortion across 8 entity types",
        5:  "Few-shot #6 — Heavy Hindi-English code-switching",
        6:  "Few-shot #7 — Cross-sentence coreference (pronouns NOT tagged)",
        7:  "Few-shot #8 — Japanese JSON record / CO-012 HR cluster (non-Latin)",
        8:  "Few-shot #9 — Arabic RTL / key_value_pairs / density=high (Article 9 + IBAN)",
        9:  "Few-shot #10 — Negation frame / ticket_worknotes (Article 9: Religion + Crime)",
        10: "Few-shot #11 — Paired-sweep mode (v3.2) / anchor A-042 / frozen entity values (R16)",
    }

    for i, fs in enumerate(data["few_shots"]):
        add_heading(doc, fs_titles.get(i, f"Few-shot #{i+1}"), level=2)
        add_heading(doc, "Axis configuration", level=3)
        axes_table(doc, fs["axes"])
        add_heading(doc, "Expected output (verbatim)", level=3)
        add_code_block(doc, pretty_json(fs["output"]))

    # ── Integration ──────────────────────────────────────────────────────────
    add_heading(doc, "4. Integration Notes", level=1)
    add_body(doc,
        "Operational notes for wiring this prompt into the SyGra graph_config.yaml. "
        "These are decisions for the pipeline, not the prompt itself.")
    integ = data.get("integration", {})
    for key, val in integ.items():
        add_heading(doc, key.replace("_", " ").title(), level=3)
        add_text_block(doc, val if isinstance(val, str) else str(val))

    # ── A→G coverage table ───────────────────────────────────────────────────
    add_heading(doc, "5. A→G Gap Coverage Map", level=1)
    add_body(doc,
        "Mapping every gap from the v3 Gap Analysis to whether and where it "
        "is addressed in this prompt. ✅ = covered in prompt; — = "
        "deliberately out of scope (sampler/pipeline/audit/process).")

    rows = [
        ("A1",  "No covering array",                             "Sampler",  "Noted in §4 sampling_strategy"),
        ("A2",  "No per-type minimum floor",                     "Sampler",  "Noted in §4 sampling_strategy"),
        ("A3",  "ML pipeline never run",                         "—",        "Out of scope"),
        ("A4",  "Indo-European/Japonic only",                    "Catalog",  "Resolved by Catalog v3 (25 langs)"),
        ("A5",  "No CheckList behavioral matrix",                "✅",       "behavioral_frame slot + few-shots #3,#4,#5,#7,#10"),
        ("B1",  "Truncation at maxTokens=8192",                  "Pipeline", "Out of scope"),
        ("B2",  "No dedup",                                      "✅",       "prior_passages_summary slot + integration.diversity_loop"),
        ("B3",  "Negation/hypothetical missing",                 "✅",       "R7 disclosed field; few-shots #3, #10"),
        ("B4",  "Partial/obfuscated missing",                    "✅",       "R7 disclosure_form, R8 OCR rules; few-shots #4, #5"),
        ("B5",  "Cross-sentence coreference",                    "✅",       "behavioral_frame=cross_sentence; few-shot #7"),
        ("C1",  "No IAA",                                        "—",        "Out of scope"),
        ("C2",  "No annotation guidelines doc",                  "✅",       "R1–R15 are de-facto guidelines"),
        ("C3",  "Nested span policy undefined",                  "✅",       "R11 triple-name rule + nesting_parent field"),
        ("C4",  "No confidence/provenance",                      "—",        "Out of scope"),
        ("C5",  "validation_dropped not audited",                "—",        "Out of scope"),
        ("D1",  "13 BLK enterprise types missing",               "Catalog",  "Resolved by Catalog v3 (51 types)"),
        ("D2",  "Digital/financial 0 patterns",                  "Catalog",  "Resolved by Catalog v3 (~820 patterns)"),
        ("D3",  "GDPR Article 9 / sensitivity",                  "✅",       "R14 sensitivity_tier; few-shots #9, #10"),
        ("E1",  "No LLM retry",                                  "—",        "Out of scope"),
        ("E2",  "Repair Claude-only",                            "—",        "Out of scope"),
        ("E3",  "No orchestration",                              "—",        "Out of scope"),
        ("E4",  "very_large+high+hard never tested",             "✅",       "R10(e) density check, R12 format-validity, R15 script-aware"),
        ("F1",  "No entropy/diversity metrics",                  "—",        "Out of scope"),
        ("F2",  "Format/checksum compliance",                    "✅",       "R12 format-validity contract (Luhn, IBAN mod-97, locale dates/numbers)"),
        ("F3",  "Per-language audit",                            "—",        "Out of scope"),
        ("F4",  "No MAUVE",                                      "—",        "Out of scope"),
        ("G1",  "No annotation guidelines doc",                  "✅",       "R1–R15 export-ready"),
        ("G2",  "No Track 2/3/4 eval scripts",                   "—",        "Out of scope (R14 enables them)"),
        ("G3",  "No corpus version registry",                    "—",        "Out of scope"),
    ]

    table = doc.add_table(rows=len(rows) + 1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    headers = ["#", "Gap", "Status", "Where in prompt"]
    h = table.rows[0]
    for i, txt in enumerate(headers):
        c = h.cells[i]
        set_cell_bg(c, ACCENT_BLUE)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        run = c.paragraphs[0].add_run(txt)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = DARK_BLUE

    for ri, row in enumerate(rows, start=1):
        rrow = table.rows[ri]
        # color rows by status
        status = row[2]
        if status == "✅":
            bg = RGBColor(0xE2, 0xEF, 0xDA)
        elif status in ("—",):
            bg = LIGHT_GREY if ri % 2 == 0 else WHITE
        elif status in ("Sampler", "Pipeline"):
            bg = RGBColor(0xFF, 0xF2, 0xCC)
        else:  # Catalog
            bg = RGBColor(0xDD, 0xEB, 0xF7)
        for ci, txt in enumerate(row):
            c = rrow.cells[ci]
            set_cell_bg(c, bg)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            run = c.paragraphs[0].add_run(str(txt))
            run.font.size = Pt(8.5)
            run.font.color.rgb = BLACK
            if ci == 2 and txt == "✅":
                run.bold = True
                run.font.color.rgb = GREEN

    # column widths
    widths = [0.45, 2.5, 0.85, 3.4]
    for i, w in enumerate(widths):
        for r in table.rows:
            r.cells[i].width = Inches(w)

    doc.add_paragraph()
    add_body(doc,
        "Score: 7/7 prompt-addressable gaps covered (A5, B3, B4, B5, C3, "
        "D3, F2). 3 catalog-level gaps resolved by Catalog v3 (A4, D1, D2). "
        "17 items correctly out of scope (sampler, pipeline, audit, process).")

    doc.save(DST)
    print(f"Wrote {DST}")
    print(f"Size: {DST.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
